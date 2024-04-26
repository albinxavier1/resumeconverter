import streamlit as st
import fitz # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH # Import the enumeration
from io import BytesIO

# Function to convert PDF to DOCX
def pdf_to_docx(pdf_path):
    doc = Document()
    with fitz.open(pdf_path) as pdf:
        for page in pdf:
            text = page.get_text("text")
            doc.add_paragraph(text)
    return doc

# Function to adjust alignment using WD_ALIGN_PARAGRAPH enumeration
def adjust_alignment(doc, alignment):
    for paragraph in doc.paragraphs:
        if alignment == "LEFT":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "CENTER":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "RIGHT":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc

# Streamlit app
st.title("Resume Converter")

uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])
if uploaded_file is not None:
    pdf_path = "temp.pdf"
    with open(pdf_path, "wb") as f:
        f.write(uploaded_file.getvalue())
    
    doc = pdf_to_docx(pdf_path)
    alignment = st.selectbox("Select Alignment", ["LEFT", "CENTER", "RIGHT"])
    doc = adjust_alignment(doc, alignment)
    
    # Save the DOCX file
    docx_path = "converted_resume.docx"
    doc.save(docx_path)
    
    # Display the converted DOCX file
    with open(docx_path, "rb") as f:
        st.download_button("Download Converted Resume", f, file_name="converted_resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
